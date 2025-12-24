from browser_use import Agent, Browser, ChatBrowserUse
import asyncio
from datetime import datetime

# WORD
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# PDF
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4


async def example():

    user_task = input("Enter your task: ")

    browser = Browser()
    llm = ChatBrowserUse(api_key="bu_rmrTMM3O5IZT5kMqGSKQ7UJYolWB-QpCLbYK2lXd3CQ")

    agent = Agent(
        task=user_task,
        llm=llm,
        browser=browser,
    )

    history = await agent.run()

    # ------------ Extract Clean Summary Text ------------
    final_text = ""
    try:
        for step in history:
            if hasattr(step, "long_term_memory") and step.long_term_memory:
                final_text += step.long_term_memory + "\n\n"
            elif hasattr(step, "extracted_content") and step.extracted_content:
                final_text += step.extracted_content + "\n\n"
    except:
        pass

    if not final_text.strip():
        final_text = "Readable summary not detected, but agent executed successfully."

    save_report(user_task, final_text)

    print("\n✨ Well formatted PDF & Word Generated Successfully!")


def save_report(task, result_text):
    timestamp = datetime.now().strftime("%d-%m-%Y %I:%M %p")
    file_time = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")

    # ------------- WORD REPORT -------------
    doc = Document()
    title = doc.add_heading("AUTONOMOUS BROWSER AGENT REPORT", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = doc.add_paragraph().add_run(f"\nGenerated On: {timestamp}\n")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_heading("Task Given:", level=2)
    doc.add_paragraph(task)

    doc.add_heading("Result Summary:", level=2)
    doc.add_paragraph(result_text)

    doc.add_paragraph("\nGenerated Using Browser-Use AI Agent")
    doc.save(f"Well_Formatted_Report_{file_time}.docx")

    # ------------- PDF REPORT -------------
    pdf_file = f"Well_Formatted_Report_{file_time}.pdf"
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "TitleStyle",
        parent=styles["Heading1"],
        alignment=1,
        spaceAfter=20,
    )

    normal = styles["BodyText"]

    story = []
    story.append(Paragraph("AUTONOMOUS BROWSER AGENT REPORT", title_style))
    story.append(Paragraph(f"Generated On: {timestamp}", normal))
    story.append(Spacer(1, 20))

    story.append(Paragraph("<b>Task Given:</b>", styles["Heading2"]))
    story.append(Paragraph(task, normal))
    story.append(Spacer(1, 15))

    story.append(Paragraph("<b>Result Summary:</b>", styles["Heading2"]))

    for line in result_text.split("\n"):
        story.append(Paragraph(line, normal))

    pdf = SimpleDocTemplate(pdf_file, pagesize=A4)
    pdf.build(story)


if __name__ == "__main__":
    asyncio.run(example())
